import os
from pathlib import Path
from typing import Dict, Any, Tuple
from docx import Document
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF not available, using PyPDF2 only")

import PyPDF2  # Fallback for PyMuPDF
from dataclasses import dataclass


@dataclass
class DocumentContent:
    text: str
    metadata: Dict[str, Any]
    file_type: str


class DocumentProcessor:
    """Processes various document formats and extracts text with metadata"""
    
    SUPPORTED_FORMATS = {
        'application/pdf': '_process_pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '_process_docx',
        'text/plain': '_process_txt'
    }
    
    def __init__(self):
        pass
    
    def process_document(self, file_path: str) -> DocumentContent:
        """Process document and extract text with metadata"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_type = self._detect_file_type_by_extension(file_path)
        
        if file_type not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        processor_method = getattr(self, self.SUPPORTED_FORMATS[file_type])
        text, metadata = processor_method(file_path)
        
        base_metadata = self._extract_base_metadata(file_path)
        metadata.update(base_metadata)
        
        return DocumentContent(
            text=text,
            metadata=metadata,
            file_type=file_type
        )
    
    def _detect_file_type_by_extension(self, file_path: str) -> str:
        """Detect file type by extension"""
        extension = Path(file_path).suffix.lower()
        
        extension_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.doc': 'application/msword',
            '.rtf': 'application/rtf',
            '.odt': 'application/vnd.oasis.opendocument.text'
        }
        
        return extension_map.get(extension, 'unknown')
    
    def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF files with robust engine testing"""
        # Normalize path for Windows
        normalized_path = os.path.normpath(file_path)
        print(f"📄 Processing PDF: {normalized_path}")
        
        # Test file readability first
        try:
            with open(normalized_path, 'rb') as test_file:
                test_file.read(1024)  # Read first 1KB to test
            print(f"✅ PDF file is readable")
        except Exception as e:
            raise ValueError(f"Cannot read PDF file: {str(e)}")
        
        # Try PyMuPDF first (more robust)
        if PYMUPDF_AVAILABLE:
            print(f"🔍 Trying PyMuPDF (fitz)...")
            try:
                return self._process_pdf_with_pymupdf(normalized_path)
            except Exception as pymupdf_error:
                print(f"❌ PyMuPDF failed: {pymupdf_error}")
        else:
            print(f"⚠️ PyMuPDF not available, skipping to PyPDF2...")
        
        # Fallback to PyPDF2
        print(f"🔄 Using PyPDF2...")
        try:
            return self._process_pdf_with_pypdf2(normalized_path)
        except Exception as pypdf2_error:
            print(f"❌ PyPDF2 also failed: {pypdf2_error}")
            
            # Both engines failed - check if it's an image-based PDF
            error_msg = f"Failed to extract text from PDF:\n"
            if PYMUPDF_AVAILABLE:
                error_msg += f"  - PyMuPDF: Failed\n"
            error_msg += f"  - PyPDF2: {str(pypdf2_error)}\n"
            error_msg += f"This PDF might be image-based and require OCR processing."
            
            raise ValueError(error_msg)
    
    def _process_pdf_with_pymupdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text using PyMuPDF with enhanced error handling"""
        text_content = []
        metadata = {'pages': 0, 'processing_engine': 'PyMuPDF'}
        
        pdf_document = fitz.open(file_path)
        metadata['pages'] = len(pdf_document)
        
        if metadata['pages'] == 0:
            pdf_document.close()
            raise ValueError("PDF has 0 pages")
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()
            
            if page_text.strip():
                text_content.append(page_text)
        
        # Extract PDF metadata
        pdf_metadata = pdf_document.metadata
        if pdf_metadata:
            metadata.update({
                'title': pdf_metadata.get('title', ''),
                'author': pdf_metadata.get('author', ''),
                'subject': pdf_metadata.get('subject', ''),
                'creator': pdf_metadata.get('creator', ''),
                'producer': pdf_metadata.get('producer', ''),
                'creation_date': pdf_metadata.get('creationDate', ''),
                'modification_date': pdf_metadata.get('modDate', '')
            })
        
        pdf_document.close()
        
        total_text = '\n\n'.join(text_content)
        
        if not total_text.strip():
            raise ValueError("No readable text found in PDF. The PDF appears to be image-based.")
        
        return total_text, metadata
    
    def _process_pdf_with_pypdf2(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text using PyPDF2 with enhanced error handling"""
        text_content = []
        metadata = {'pages': 0, 'processing_engine': 'PyPDF2'}
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata['pages'] = len(pdf_reader.pages)
            
            if metadata['pages'] == 0:
                raise ValueError("PDF has 0 pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception:
                    pass
            
            # Extract PDF metadata
            if pdf_reader.metadata:
                pdf_meta = pdf_reader.metadata
                metadata.update({
                    'title': pdf_meta.get('/Title', ''),
                    'author': pdf_meta.get('/Author', ''),
                    'subject': pdf_meta.get('/Subject', ''),
                    'creator': pdf_meta.get('/Creator', '')
                })
        
        total_text = '\n\n'.join(text_content)
        
        if not total_text.strip():
            raise ValueError("No readable text found in PDF. The PDF appears to be image-based.")
        
        return total_text, metadata
    
    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX files"""
        doc = Document(file_path)
        
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'subject': doc.core_properties.subject or '',
            'keywords': doc.core_properties.keywords or ''
        }
        
        return '\n\n'.join(text_content), metadata
    
    def _process_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        metadata = {
            'lines': len(content.splitlines()),
            'characters': len(content)
        }
        
        return content, metadata
    
    def _extract_base_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic file metadata"""
        path_obj = Path(file_path)
        stat = path_obj.stat()
        
        return {
            'filename': path_obj.name,
            'file_size': stat.st_size,
            'created_at': stat.st_ctime,
            'modified_at': stat.st_mtime,
            'file_extension': path_obj.suffix.lower()
        } 